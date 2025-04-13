"""
utils.py

Contains utility/helper functions for file processing, text extraction,
filename sanitization, contact info parsing, and generating .docx/.pdf/.zip files.
"""


import os
import io
import re
import requests
from docx import Document
import pdfplumber
import zipfile
from bs4 import BeautifulSoup
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas


def safe_filename(file):
    return secure_filename(file.filename)


def extract_text_from_file(file):
    """
    Extracts raw text content from a file (.txt, .pdf, or .docx).

    Args:
        file (FileStorage): Uploaded file object from the form.

    Returns:
        str: Extracted text content.
    """
    filename = safe_filename(file)
    ext = os.path.splitext(filename)[1].lower()  # Get file extension to determine parsing logic

    if ext == ".txt":
        return file.read().decode("utf-8")
    elif ext == ".pdf":
        with pdfplumber.open(file) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif ext == ".docx":
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        raise ValueError("Unsupported file type")
    
def extract_contact_info(job_description):
    email_match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", job_description)
    phone_match = re.search(r"\+?\d[\d\s\-()]{7,}\d", job_description)

    contact = []
    if email_match:
        contact.append(f"Email: {email_match.group()}")
    if phone_match:
        contact.append(f"Phone: {phone_match.group()}")

    return " | ".join(contact) if contact else "Not provided"

def generate_docx_file(text):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_pdf_file(text):
    """
    Generates a PDF file stream from the given plain text using reportlab.

    Args:
        text (str): Text content to convert into PDF.

    Returns:
        BytesIO: File-like object representing the PDF file.
    """
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER

    # Split and draw each line
    y = height - 50 # Starting Y position for text drawing on the PDF page
    for line in text.splitlines():
        c.drawString(50, y, line)
        y -= 15
        if y < 50:
            c.showPage()
            y = height - 50

    c.save()
    buffer.seek(0)
    return buffer

def create_zip(files: dict, zip_name="output.zip"):
    """
    Creates a ZIP archive in memory from a dictionary of file names and string content.

    Args:
        files (dict): Dictionary where keys are filenames and values are text content.

    Returns:
        BytesIO: A file-like ZIP stream ready for download.
    """

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for filename, content in files.items():
            zip_file.writestr(filename, content)
    zip_buffer.seek(0)
    return zip_buffer

def extract_job_description_from_url(job_url):

    try:
        response = requests.get(job_url, timeout=5)
        response.raise_for_status()  # Raises HTTPError for bad responses
        soup = BeautifulSoup(response.text, "html.parser")

        # Extract both paragraphs and list items
        paragraphs = soup.find_all(["p", "li"])
        job_description = "\n".join(p.get_text() for p in paragraphs).strip()
        return job_description if job_description else None
    except Exception:
        return None