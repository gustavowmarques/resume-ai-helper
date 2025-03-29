from flask import Flask, render_template, request, session, redirect, url_for, send_file
from dotenv import load_dotenv
from flask_session import Session 
import os 
from openai import OpenAI

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

import requests
from bs4 import BeautifulSoup
import docx
import pdfplumber
from werkzeug.utils import secure_filename

from docx import Document
import io

import smtplib
from email.message import EmailMessage

import zipfile

app = Flask(__name__)
app.secret_key = "supersecretkey"

app.config["SESSION_TYPE"] = "filesystem"
Session(app)

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/upload_resume", methods=["GET", "POST"])
def upload_resume():
    print(" /upload_resume triggered")

    if request.method == "POST":
        print("POST method received on /upload_resume")

        resume_text = request.form.get("resume", "").strip()
        print(f"Pasted resume length: {len(resume_text)} characters")

        uploaded_file = request.files.get("resume_file")

        if uploaded_file and uploaded_file.filename != "":
            filename = secure_filename(uploaded_file.filename)
            file_ext = os.path.splitext(filename)[1].lower()
            print(f"Uploaded file: {filename}")

            try:
                if file_ext == ".txt":
                    resume_text = uploaded_file.read().decode("utf-8")

                elif file_ext == ".pdf":
                    with pdfplumber.open(uploaded_file) as pdf:
                        resume_text = "\n".join(page.extract_text() or "" for page in pdf.pages)

                elif file_ext == ".docx":
                    doc = docx.Document(uploaded_file)
                    resume_text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    return render_template("upload_resume.html", error="Unsupported file format.")

                print(f"Extracted resume from file: {len(resume_text)} characters")

            except Exception as e:
                print(f"Error reading file: {e}")
                return render_template("upload_resume.html", error="There was an error processing the uploaded file.")

        if not resume_text:
            return render_template("upload_resume.html", error="Please upload a resume file or paste the text.")

        
        session["resume"] = resume_text
        return redirect(url_for("job_description"))

    # Default GET request
    return render_template("upload_resume.html")


@app.route("/test_form", methods=["GET", "POST"])
def test_form():
    if request.method == "POST":
        print("✅ Test form POST received")
        return "Success!"

    return '''
    <form method="POST">
      <textarea name="resume"></textarea>
      <button type="submit">Submit</button>
    </form>
    '''

@app.route("/job_description", methods=["GET", "POST"])
def job_description():
    resume = session.get("resume", "") 

    if request.method == "POST":
        job_desc = request.form.get("job", "")
        job_url = request.form.get("job_url", "")

        if not job_desc and not job_url:
            return render_template("job_description.html", resume=resume, error="Please paste a job description or enter a job URL.")

        if job_url:
            try:
                response = requests.get(job_url, timeout=5)
                soup = BeautifulSoup(response.text, "html.parser")
                paragraphs = soup.find_all(["p", "li"])
                job_desc = "\n".join(p.get_text() for p in paragraphs).strip()

                if not job_desc:
                    return render_template("job_description.html", resume=resume, error="Unable to extract job description from this URL. Try copying and pasting it instead.")
                
            except Exception as e:
                return render_template("job_description.html", resume=resume, error=f"Error fetching job description: {str(e)}")

        session["job"] = job_desc
        session["job_url"] = job_url

        return redirect(url_for("ai_suggestions"))

    return render_template("job_description.html", resume=resume,
                       job=session.get("job", ""),
                       job_url=session.get("job_url", ""))


@app.route("/cover_letter", methods=["POST"])
def cover_letter():
    resume = request.form.get("resume", "")
    job = request.form.get("job", "")

    user_name = request.form.get("name", "Your Name")
    job_title = request.form.get("job_title", "the advertised position")
    tone = request.form.get("tone", "formal")
    language = request.form.get("language", "English")

    prompt = f"""
Using the resume and job description below, write a personalized cover letter for a position titled '{job_title}' in {language}.

Resume:
{resume}

Job Description:
{job}

Use a {tone} tone in your writing. Sign it off with: {user_name}
"""

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=600,
            temperature=0.7
        )

        generated_letter = response.choices[0].message.content.strip()
        

    except Exception as e:
        generated_letter = f"OpenAI API Error: {str(e)}"
        
    session["cover_letter"] = generated_letter

    print("Generated cover letter:", generated_letter[:100])
    print("Stored in session:", len(session.get("cover_letter", "")))


    return render_template(
        "cover_letter.html",
        letter=generated_letter,
        resume=resume,
        job=job,
        tone=tone,
        name=user_name,
        job_title=job_title,
        response=session.get("ai_suggestions", "")
        )

@app.route("/ai_suggestions", methods=["GET"])
def ai_suggestions():
    resume = session.get("resume", "")
    lines = resume.splitlines()
    name_guess = lines[0].strip() if lines else ""

    job = session.get("job", "")
    job_title_guess = ""

    if job:
        job_title_guess = job.splitlines()[0].strip()
    elif "job_url" in session:
        job_url = session["job_url"]
        job_title_guess = job_url.split("/")[-1].replace("-", " ").replace("_", " ").split("?")[0].title()

    if not resume or not job:
        return render_template(
            "ai_suggestions.html",
            response="Missing resume or job description.",
            resume="",
            job="",
            name_guess=name_guess)

    resume = resume[:3000]
    job = job[:2000]

    prompt = f"""
You are an expert career coach. Compare the following resume with the job description and provide personalized, actionable suggestions to improve the resume for this role.

Resume:
{resume}

Job Description:
{job}

Suggestions:
"""

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=500,
            temperature=0.7
        )

        ai_response = response.choices[0].message.content.strip()
        session["ai_suggestions"] = ai_response


    except Exception as e:
        ai_response = f" OpenAI API Error: {str(e)}"

    return render_template(
        "ai_suggestions.html",
        response=ai_response, 
        resume=resume,
        job=job,
        name_guess=name_guess,
        job_title_guess=job_title_guess,
        )

@app.route("/download_docx")
def download_docx():
    letter = session.get("cover_letter", "")

    if not letter:
        return "No cover letter to download."

    doc = Document()
    for line in letter.split('\n'):
        doc.add_paragraph(line)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="cover_letter.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/feedback", methods=["GET", "POST"])
def feedback():
    if request.method == "POST":
        feedback_text = request.form.get("feedback", "").strip()
        email = request.form.get("email", "").strip()

        print("Feedback received:", feedback_text)
        print("Optional email:", email)

        # Save to file
        with open("feedback.txt", "a", encoding="utf-8") as f:
            f.write("----- Feedback Entry -----\n")
            if email:
                f.write(f"Email: {email}\n")
            f.write(f"Feedback: {feedback_text}\n")
            f.write("--------------------------\n\n")

        # SEND EMAIL HERE — INSIDE POST BLOCK
        gmail_user = os.getenv("MAIL_USERNAME")
        gmail_password = os.getenv("MAIL_PASSWORD")

        msg = EmailMessage()
        msg["Subject"] = "New Feedback from Resume AI Helper"
        msg["From"] = gmail_user
        msg["To"] = gmail_user  # You can change this to any destination

        body = f"Feedback:\n{feedback_text}\n\nEmail: {email or 'Not provided'}"
        msg.set_content(body)

        try:
            with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
                smtp.login(gmail_user, gmail_password)
                smtp.send_message(msg)
                print("Feedback email sent.")
        except Exception as e:
            print(f"Failed to send feedback email: {e}")

        return render_template("feedback_thankyou.html")

    return render_template("feedback.html")

@app.errorhandler(404)
def page_not_found(e):
    return render_template("404.html"), 404

@app.errorhandler(500)
def internal_server_error(e):
    return render_template("500.html"), 500

@app.route("/download_zip")
def download_zip():
    # Get content from session
    suggestions = session.get("ai_suggestions", "")
    cover_letter = session.get("cover_letter", "")

    if not suggestions and not cover_letter:
        return "No content to zip", 400

    # Create in-memory ZIP file
    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, "w") as zf:
        if suggestions:
            zf.writestr("ai_suggestions.txt", suggestions)
        if cover_letter:
            zf.writestr("cover_letter.txt", cover_letter)

    memory_file.seek(0)

    return send_file(
        memory_file,
        download_name="resume_ai_helper.zip",
        as_attachment=True
    )

@app.route("/start_over")
def start_over():
    session.clear()
    return redirect(url_for("home"))

if __name__ =="__main__":
    app.run(debug=True)